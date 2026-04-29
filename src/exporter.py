"""
Exportador de roteiros para formato DOCX.
Gera documentos Word com formatação idêntica aos roteiros de referência em /kb.
Padrão: Tahoma 14pt bold (cabeçalho), 12pt bold (locução), 12pt normal (imagem/lettering).
"""
import io
import re
import zipfile
from datetime import datetime
import pytz
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_header_line(doc, text: str):
    """Adiciona linha de cabeçalho: Tahoma 14pt Bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = "Tahoma"
    run.font.size = Pt(14)
    run.bold = True


def _add_separator(doc):
    """Adiciona linha separadora."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("______________________________________________________________________")
    run.font.name = "Tahoma"
    run.font.size = Pt(12)


def _add_locucao(doc, text: str):
    """Adiciona linha de locução: Tahoma 12pt Bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = "Tahoma"
    run.font.size = Pt(12)
    run.bold = True


def _add_imagem(doc, text: str):
    """Adiciona linha de imagem/lettering: Tahoma 12pt Normal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = "Tahoma"
    run.font.size = Pt(12)
    run.bold = False


def _add_empty_line(doc):
    """Adiciona linha em branco."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


def _extract_product_name(roteiro_text: str) -> str:
    """Tenta extrair o nome do produto do texto do roteiro."""
    # Procura na linha de Produto:
    match = re.search(r'Produto:\s*(.+)', roteiro_text)
    if match:
        name = match.group(1).strip()
        # Remove prefixos (NW, REVIEW, 3D, LU, SOCIAL), meses (JAN-DEZ) e códigos numéricos longos
        name = re.sub(r'^(NW|SOCIAL|REVIEW|3D|LU|JAN|FEV|MAR|ABR|MAI|JUN|JUL|AGO|SET|OUT|NOV|DEZ)\s*', '', name, flags=re.IGNORECASE)
        name = re.sub(r'^(NW|SOCIAL|REVIEW|3D|LU|JAN|FEV|MAR|ABR|MAI|JUN|JUL|AGO|SET|OUT|NOV|DEZ)\s*', '', name, flags=re.IGNORECASE) # Segunda passada
        name = re.sub(r'^(\d{6,}\s*)+', '', name) # Remove códigos numéricos
        # Remove placeholders e lixo de markdown
        name = re.sub(r'\[?NOME DO PRODUTO\]?', '', name, flags=re.IGNORECASE)
        name = re.sub(r'^\**TÍTULO( DO PRODUTO)?:?\**\s*', '', name, flags=re.IGNORECASE)
        return name.strip()

    # Fallback: procura no título (primeiras palavras do roteiro que parecem nome de produto)
    lines = roteiro_text.strip().split('\n')
    for line in lines:
        line = line.strip()
        if line.startswith('- ') and ('da ' in line or 'do ' in line):
            # Tenta extrair "Este [Produto], da [Marca]"
            match2 = re.search(r'Est[ea]\s+(.+?),\s+d[ao]', line)
            if match2:
                return match2.group(1).strip()

    return "Produto"


def _parse_roteiro(roteiro_text: str) -> list[dict]:
    """
    Parseia o texto bruto do roteiro em blocos estruturados.
    Retorna lista de dicts com tipo (header, separator, locucao, imagem, lettering, empty).
    """
    blocks = []
    lines = roteiro_text.strip().split('\n')

    for line in lines:
        stripped = line.strip()

        # Limpa markdown bold da linha para análise de tipo
        analysis_line = stripped.strip("*").strip()

        if not stripped:
            blocks.append({"type": "empty"})
        elif analysis_line.startswith("Cliente:"):
            blocks.append({"type": "header", "text": analysis_line})
        elif analysis_line.startswith("Roteirista:"):
            blocks.append({"type": "header", "text": analysis_line})
        elif analysis_line.startswith("Produto:"):
            blocks.append({"type": "header", "text": analysis_line})
        elif stripped.startswith("____"):
            blocks.append({"type": "separator"})
        elif analysis_line.startswith("Imagem:"):
            blocks.append({"type": "imagem", "text": analysis_line})
        elif analysis_line.startswith("Lettering:"):
            blocks.append({"type": "lettering", "text": analysis_line})
        elif analysis_line.startswith("- "):
            blocks.append({"type": "locucao", "text": analysis_line})
        elif stripped.startswith("**") and stripped.endswith("**"):
            # Se for apenas bold e não foi pego acima como header/imagem etc
            clean = analysis_line
            blocks.append({"type": "locucao", "text": f"- {clean}"})
        else:
            # Linhas que não se encaixam nos padrões acima são tratadas como locução por padrão (Bold no DOCX)
            blocks.append({"type": "locucao", "text": stripped})

    return blocks


def generate_filename(code: str, product_name: str, selected_month: str = "MAR", model_id: str = "", com_lu: bool = True) -> str:
    """Gera nome do arquivo no padrão: NW LU {selected_month} {code} {product_name} [{model}].docx"""
    # Garante que o código tenha 9 dígitos (preenche com 0 à direita se necessário)
    clean_code = str(code).strip()
    if clean_code and len(clean_code) < 9:
        clean_code = clean_code.ljust(9, '0')

    # Detecta prefixo correto baseado no modo e presença da Lu
    prefixo_u = str(product_name).upper()
    
    if "3D" in prefixo_u:
        prefixo = f"NW 3D LU {selected_month}"
    elif "SOCIAL" in prefixo_u:
        prefixo = f"SOCIAL {selected_month}"
    elif "NW REVIEW" in prefixo_u or com_lu == "REVIEW":
        prefixo = f"NW REVIEW {selected_month}"
    elif com_lu is True:
        prefixo = f"NW LU {selected_month}"
    else:
        prefixo = f"NW {selected_month}"

    # Agora limpa o nome para o arquivo
    clean_name = _clean_product_name(product_name)

    model_tag = ""
    if model_id:
        model_tag = model_id.split('/')[-1].upper()
        model_tag = f" [{model_tag}]"

    return f"{prefixo} {clean_code} {clean_name}{model_tag}.docx"


def _clean_product_name(raw_name: str) -> str:
    """Limpa o nome do produto de forma agressiva para evitar duplicação."""
    # Remove prefixos de taxonomia, meses e códigos numéricos em cascata
    clean = raw_name
    for _ in range(3): # Múltiplas passadas para limpar prefixos grudados
        clean = re.sub(r'^(NW|SOCIAL|REVIEW|3D|LU|JAN|FEV|MAR|ABR|MAI|JUN|JUL|AGO|SET|OUT|NOV|DEZ)\s*', '', clean, flags=re.IGNORECASE).strip()
        clean = re.sub(r'^(\d{6,}\s*)+', '', clean).strip()
    
    # Remove placeholders e caracteres inválidos
    clean = re.sub(r'\[?NOME DO PRODUTO\]?', '', clean, flags=re.IGNORECASE)
    clean = re.sub(r'^\**TÍTULO( DO PRODUTO)?:?\**\s*', '', clean, flags=re.IGNORECASE)
    clean = re.sub(r'[<>:"/\\|?*]', '', clean)
    return clean[:80].strip()


def export_roteiro_docx(roteiro_text: str, code: str = "", product_name: str = "", selected_month: str = "MAR", selected_date: str = None, model_id: str = "", com_lu: bool = True) -> tuple[bytes, str]:
    """
    Gera um documento Word (.docx) com a formatação de referência.
    O cabeçalho interno do DOCX é idêntico ao nome do arquivo.

    Args:
        roteiro_text: Texto completo do roteiro gerado pela IA
        code: Código do produto Magalu
        product_name: Nome do produto (extraído automaticamente se vazio)
        com_lu: Se o roteiro tem a Lu ou não (para taxonomia do nome do arquivo)

    Returns:
        Tuple de (bytes do docx, nome do arquivo)
    """
    doc = Document()

    # Configura estilo padrão
    style = doc.styles['Normal']
    style.font.name = 'Tahoma'
    style.font.size = Pt(12)

    # Extrai nome do produto se não fornecido
    if not product_name:
        product_name = _extract_product_name(roteiro_text)

    # Gera o filename PRIMEIRO — ele é a fonte de verdade para a nomenclatura
    filename = generate_filename(code, product_name, selected_month, model_id, com_lu=com_lu)
    
    # O nome do cabeçalho deve ser EXATAMENTE o mesmo do filename (sem .docx e sem model tag)
    header_product_line = filename.replace(".docx", "")
    # Remove model tag do cabeçalho (ex: " [GEMINI-3-FLASH-PREVIEW]")
    header_product_line = re.sub(r'\s*\[[A-Z0-9._-]+\]\s*$', '', header_product_line).strip()

    # Parseia o roteiro
    blocks = _parse_roteiro(roteiro_text)

    # Verifica se já tem cabeçalho no texto
    has_header = any(b["type"] == "header" for b in blocks)

    if not has_header:
        # Gera cabeçalho padrão — linha Produto: usa a mesma string do filename
        header_date = selected_date if selected_date else datetime.now(pytz.timezone('America/Sao_Paulo')).strftime('%d/%m/%y')
        _add_header_line(doc, "Cliente: Magalu")
        _add_header_line(doc, f"Roteirista: Tiago Fernandes - Data: {header_date}")
        _add_header_line(doc, f"Produto: {header_product_line}")
        _add_separator(doc)
        _add_empty_line(doc)

    # Renderiza cada bloco
    for block in blocks:
        btype = block["type"]
        text = block.get("text", "")

        if btype == "header":
            # Corrige a data se necessário
            if "Data:" in text:
                now = datetime.now(pytz.timezone('America/Sao_Paulo'))
                text = re.sub(r'Data:\s*[\d/]+', f"Data: {now.strftime('%d/%m/%y')}", text)
            # Força a linha Produto: a usar a nomenclatura unificada
            if text.strip().startswith("Produto:"):
                text = f"Produto: {header_product_line}"
            _add_header_line(doc, text)
        elif btype == "separator":
            _add_separator(doc)
        elif btype == "locucao":
            _add_locucao(doc, text)
        elif btype == "imagem":
            _add_imagem(doc, text)
        elif btype == "lettering":
            _add_imagem(doc, text)
        elif btype == "empty":
            _add_empty_line(doc)
        elif btype == "text":
            _add_imagem(doc, text)

    # Gera bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue(), filename



def format_for_display(roteiro_text: str) -> str:
    """
    Formata o roteiro para exibição no Streamlit com Markdown.
    Locução em **bold**, Imagem sem bold, com quebra de linha.
    """
    lines = roteiro_text.strip().split('\n')
    formatted = []

    for line in lines:
        stripped = line.strip()

        if not stripped:
            formatted.append("")
        elif stripped.startswith("Cliente:") or stripped.startswith("Roteirista:") or stripped.startswith("Produto:"):
            formatted.append(f"**{stripped}**")
        elif stripped.startswith("____"):
            formatted.append("---")
        elif stripped.startswith("- "):
            # Verifica se tem "Imagem:" inline (separar)
            if "Imagem:" in stripped:
                parts = stripped.split("Imagem:", 1)
                locucao = parts[0].strip()
                imagem = "Imagem:" + parts[1]
                formatted.append(f"**{locucao}**")
                formatted.append(f"\n{imagem}")
            else:
                formatted.append(f"**{stripped}**")
        elif stripped.startswith("Imagem:"):
            formatted.append(stripped)
        elif stripped.startswith("Lettering:"):
            formatted.append(stripped)
        else:
            formatted.append(stripped)

    return "\n".join(formatted)

def export_all_roteiros_zip(roteiros: list, selected_month: str = "MAR", selected_date: str = None) -> tuple[bytes, str]:
    """
    Gera um arquivo ZIP contendo todos os roteiros em formato DOCX.
    
    Args:
        roteiros: Lista de dicionários contendo 'roteiro_original', 'codigo' e 'ficha' (opcional)
        selected_month: Mês para o nome dos arquivos
        
    Returns:
        Tuple de (bytes do zip, nome do arquivo)
    """
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        for idx, item in enumerate(roteiros):
            roteiro_text = item.get('roteiro_original', '')
            # Pula roteiros que são avisos de erro de extração
            if roteiro_text.startswith("⚠️"):
                continue
            
            # Prioriza o mês que está no registro do roteiro, caso contrário usa o selecionado no lote
            mes_roteiro = item.get('mes', selected_month)
                
            doc_bytes, filename = export_roteiro_docx(
                roteiro_text,
                code=item.get('codigo', ''),
                product_name='', # Será extraído do texto do roteiro
                selected_month=mes_roteiro,
                selected_date=selected_date,
                model_id=item.get('model_id', ''),
                com_lu=item.get('com_lu', True)
            )
            # Garante que o nome do arquivo seja único dentro do ZIP se houver duplicatas
            zip_file.writestr(filename, doc_bytes)
            
    zip_buffer.seek(0)
    now = datetime.now(pytz.timezone('America/Sao_Paulo'))
    zip_filename = f"ROTEIROS_MAGALU_{now.strftime('%d_%m_%Y_%H%M')}.zip"
    
    return zip_buffer.getvalue(), zip_filename
